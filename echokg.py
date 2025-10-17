import pandas as pd
import numpy as np
import docx
import streamlit as st
from streamlit.logger import get_logger
from datetime import date
from datetime import datetime
from keras.models import load_model
from tensorflow import keras

# Загрузка таблиц и моделей
model_loaded = load_model('models/af_predictor_model_4.h5')
data = pd.read_excel('tabs/exams.xlsx')
multiselect = pd.read_excel('tabs/multiselect.xlsx')

# Функция нейросети
def prediction(ao, la, lvd, rv, ra_x, pa, ivs, pw, ef, ar, mr, tr, cv):    
    ao1 = (ao - 1.8) / (5.4 - 1.8)
    la1 = (la - 2.4) / (7.8 - 2.4)
    lvd1 = (lvd - 2.4) / (8.1 - 2.4)
    rv1 = (rv - 1.5) / (4.8 - 1.5)
    ra_x1 = (ra_x - 2.3) / (6.0 - 2.3)
    pa1 = (pa - 1.0) / (3.6 - 1.0)
    ivs1 = (ivs - 0.7) / (2.6 - 0.7)
    pw1 = (pw - 0.7) / (2.2 - 0.7)
    ef1 = (lvd - 18.0) / (75.0 - 18.0)       # должна быть ef
    ar1 = (ar - 0) / (2.5 - 0)
    mr1 = (mr - 0) / (3.5 - 0)
    tr1 = (tr - 0) / (3.5 - 0)
    cv1 = (cv - 0) / (3.5 - 0)
    x = [ao1, la1, lvd1, rv1, ra_x1, pa1, ivs1, pw1, ef1, ar1, mr1, tr1, cv1]
    x = np.array(x)
    nn = np.expand_dims(x, axis=0)
    res = model_loaded.predict(nn)    
    return res

LOGGER = get_logger(__name__)
def run():
    st.set_page_config(
        page_title="Трансторакальная эхокардиография",
        page_icon="https://cdn-icons-png.freepik.com/512/2393/2393002.png"
    )
# НАЧАЛО ФУНКЦИИ АВТООБНОВЛЕНИЯ СТРАНИЦЫ

# БОКОВАЯ ПАНЕЛЬ _______________________________________________________
	
    st.sidebar.write('ВЫБЕРИТЕ ИССЛЕДОВАНИЕ ИЗ СПИСКА ИЛИ СОЗДАЙТЕ НОВОЕ')
    
    # Выбор исследования из списка
    chose_variant = False # Отключение окна выбора пациента при регистрации нового
    chosen_exam = st.sidebar.selectbox(
    'Выберите пациента и исследование',
    (data['id_of_exam']), index=None, 
    placeholder='Выберите из списка или начните вводить Ф.И.О.',
    disabled=chose_variant)
    if chosen_exam != None:
        chose_variant = True
        anti_chose_variant = False
    else:
        chose_variant = False
        anti_chose_variant = True
		
    # Выбор строки из таблицы с данными пациента
    chosen_data = data.loc[data['id_of_exam'] == chosen_exam]
	# Имя
    chosen_name = str(chosen_data['name_of_patient'].tolist())
    chosen_name = chosen_name[2:-2]
	# Дата рождения
    chosen_date_of_birth = str(chosen_data['date_of_birth'].tolist())
    chosen_date_of_birth = chosen_date_of_birth[2:-2]
    # Дата исследования
    this_day = str(date.today())
    day = this_day[-2:] + '.'
    month = this_day[-5:-3] + '.'
    year =  this_day[:4]
    today_day = day + month + year
    chosen_date_of_exam = str(chosen_data['date_of_exam'].tolist())
    chosen_date_of_exam = chosen_date_of_exam[2:-2]
    # Время исследования
    this_time = str(datetime.now())
    today_time = this_time[10:16]
    chosen_time_of_exam = str(chosen_data['time_of_exam'].tolist())
    chosen_time_of_exam = chosen_time_of_exam[2:-2]
    # Рост пациента
    try:
        chosen_height = int(chosen_data['patient_height'])
    except:
        chosen_height = 170
    # Вес пациента
    try:
        chosen_weight = int(chosen_data['patient_weight'])
    except:
        chosen_weight = 70
  
    # Регистрация нового пациента (ФИО)
    name_of_patient = st.sidebar.text_input('Введите Ф.И.О. пациента',
        value = str(chosen_name),
        placeholder = 'Фамилия Имя Отчество',
        disabled = chose_variant)
        
    # Проверка совпадений ФИО в столбце
    if name_of_patient in data['name_of_patient'].values and chosen_exam == None:
        st.sidebar.warning('Пациент с таким Ф.И.О. уже зарегистрирован в системе',
        icon="⚠️")    
    
     # Регистрация нового пациента (дата рождения)
    date_of_birth = st.sidebar.text_input('Введите дату рождения пациента',
        value = str(chosen_date_of_birth),
        placeholder = 'дд.мм.гггг',
        disabled = chose_variant)
            
    # Возможность редактировать имеющееся исследование
    check_redact_exam = st.sidebar.checkbox('РЕДАКТИРОВАТЬ ВЫБРАННОЕ ИССЛЕДОВАНИЕ',
        value = False, disabled = anti_chose_variant)
    if check_redact_exam == True:
#        check_redact_exam = False
        chose_variant = False
    
    # Возможность вставлять в исследование текущие дату и время
    check_now_date_and_time = st.sidebar.checkbox('Вставить текущие дату и время',
        value = False, disabled = chose_variant)
    if check_now_date_and_time == True:
        chosen_date_of_exam = today_day
        chosen_time_of_exam = today_time
    
    # Регистрация нового исследования
    date_of_exam = st.sidebar.text_input('Введите дату исследованния',
        value = str(chosen_date_of_exam),
        placeholder = 'дд.мм.гггг',
        disabled = chose_variant)
    time_of_exam = st.sidebar.text_input('Введите время исследованния',
        value = str(chosen_time_of_exam),
        placeholder = 'чч:чч',
        disabled = chose_variant)
    patient_height = st.sidebar.number_input('Введите рост пациента в см',
		min_value = 50, max_value = 240,
		value = chosen_height,
		disabled = chose_variant)
    patient_weight = st.sidebar.number_input('Введите вес пациента в кг',
		min_value = 3, max_value = 180,
		value = chosen_weight,
		disabled = chose_variant)
    
    calculate_body_area = round(0.007184 * (patient_weight**0.425) * (patient_height**0.725), 2)

# ОСНОВНАЯ СТРАНИЦА - ИЗМЕРЕНИЯ ________________________________________
	
	# Поиск и присвоение номера исследования
    row_index = data.index[data['id_of_exam'] == chosen_exam].tolist()
    exam_number = str(row_index)
    try:
	    exam_number = int(exam_number[1:-1]) + 1
    except:
	    exam_number = int(len(data)) + 1
    st.subheader('ПРОТОКОЛ ЭХОКАРДИОГРАФИИ № ' + str(exam_number))
    
    # Удаление исследования
    if st.sidebar.button('Удалить выбранное исследование',
                  disabled=anti_chose_variant):
#                    try:
                    data.drop(row_index, axis=0, inplace=True)
                    data.to_excel('tabs/exams.xlsx', index=False)
                    st.sidebar.success('Исследование удалено', icon="✅")
#                    except:
#                        st.sidebar.error('Удаление не произведено!', icon="⚠️")

       

    
    # Формирование ID исследования
    if name_of_patient != '':
        id_of_exam = str(name_of_patient + ' ' +  date_of_birth + ' г.р. ТТЭ от ' + date_of_exam + ' ' + time_of_exam)
    else:
	    id_of_exam = ''
    st.text_input('Данные об исследовании: ', 
			id_of_exam,
			disabled=True) 

	# Описательная часть протокола. УЗ-сканер, ритм и особенности исследования
    col1, col2 = st.columns(2)
    chosen_scanner = str(chosen_data['scanner'].tolist())
    chosen_scanner = chosen_scanner[2:-2]
    chosen_rythm = str(chosen_data['rythm'].tolist())
    chosen_rythm = chosen_rythm[2:-2]
    chosen_visualise_troubles = str(chosen_data['visualise_troubles'].tolist())
    chosen_visualise_troubles = chosen_visualise_troubles[2:-2]        
    with col1:
	    if chosen_exam == None:
             scanner = st.selectbox(
		     'УЗ-аппарат: ',
		     (multiselect['scanner']), index=0, 
		     disabled=False)
	    else:			
		     scanner = st.text_input('УЗ-аппарат: ', 
				chosen_scanner,
				disabled = chose_variant)
    with col2:
        if chosen_exam == None: 
             rythm = st.selectbox(
		     'Ритм: ',
		     (multiselect['rythm']), index=0, 
		     placeholder='Выберите основной ритм пациента',
		     disabled=chose_variant)
        else:			
            rythm = st.text_input('Ритм: ', 
				chosen_rythm,
				disabled = chose_variant)		
    if chosen_exam == None:
        visualise_troubles = st.multiselect('Причины неоптимальной визуализации: ',
	    options = (multiselect['visualise_troubles']),
	    default=(multiselect['visualise_troubles'].iloc[0]),
		placeholder='Выберите из списка',
		disabled=chose_variant)
        visualise_troubles = ', '.join(visualise_troubles)
    else:			
        visualise_troubles = st.text_input('Причины неоптимальной визуализации',
			chosen_visualise_troubles,
			disabled=chose_variant)
   
    st.write('ОСНОВНЫЕ ПОКАЗАТЕЛИ СТРУКТУРЫ И ФУНКЦИИ СЕРДЦА')
    st.write('PLAX') 
    
    if rythm == 'Ритм синусовый' and chosen_exam == None:
        dyastola = False
    elif rythm == 'Ритм синусовый' and chose_variant == False:
        dyastola = False
    else:
        dyastola = True
    
    # Состояние стенки аорты
    chosen_aorta_wall = str(chosen_data['aorta_wall'].tolist())
    chosen_aorta_wall = chosen_aorta_wall[2:-2]
    if chosen_exam == None:
        aorta_wall = st.multiselect('Стенка аорты: ',
	    options = (multiselect['aorta_wall']),
	    default=(multiselect['aorta_wall'].iloc[1]),
		disabled=chose_variant)
        aorta_wall = ', '.join(aorta_wall)
    else:			
        aorta_wall = st.text_input('Стенка аорты:', 
			value = chosen_aorta_wall,
			disabled=chose_variant)
    
    # Размеры аорты и раскрытие аортального клапана
    col1, col2, col3 = st.columns(3)
    with col1:
        try:
            chosen_ao_valsalva = float(chosen_data['ao_valsalva'])
        except:
            chosen_ao_valsalva = 0.0
        ao_valsalva = st.number_input('Аорта, с. Вальсальвы, см',
		    min_value = 0.0, max_value = 10.0,
		    value = chosen_ao_valsalva,
		    disabled = chose_variant)
    with col2:
        try:
            chosen_ao = float(chosen_data['ao'])
        except:
            chosen_ao = 0.0
        ao = st.number_input('Восходящая аорта, см',
		    min_value = 0.0, max_value = 10.0,
		    value = chosen_ao,
		    disabled = chose_variant)
    with col3:
        try:
            chosen_av_plax = float(chosen_data['av_plax'])
        except:
            chosen_av_plax = 0.0
        av_plax = st.number_input('Раскрытие АоК, см',
		    min_value = 0.0, max_value = 4.0,
		    value = chosen_av_plax,
		    disabled = chose_variant)
	
	# Левое предсердие, правый желудочек, толщина стенки правого желудочка
    col4, col5, col6 = st.columns(3)
    with col4:
        try:
            chosen_la = float(chosen_data['la'])
        except:
            chosen_la = 0.0
        la = st.number_input('Левое предсердие, см',
		    min_value = 0.0, max_value = 10.0,
		    value = chosen_la,
		    disabled = chose_variant)
    with col5:
        try:
            chosen_rv = float(chosen_data['rv'])
        except:
            chosen_rv = 0.0
        rv = st.number_input('Правый желудочек, см',
		    min_value = 0.0, max_value = 10.0,
		    value = chosen_rv,
		    disabled = chose_variant)
    with col6:
        try:
            chosen_rv_wall = float(chosen_data['rv_wall'])
        except:
            chosen_rv_wall = 0.4
        rv_wall = st.number_input('Толщина стенки ПЖ, см',
		    min_value = 0.0, max_value = 2.0,
		    value = chosen_rv_wall,
		    disabled = chose_variant)
    
    # МЖП, КДР, ЗСЛЖ
    col7, col8, col9 = st.columns(3)
    with col7:
        try:
            chosen_ivs = float(chosen_data['ivs'])
        except:
            chosen_ivs = 0.0
        ivs = st.number_input('ТМЖПд, см',
		    min_value = 0.0, max_value = 5.0,
		    value = chosen_ivs,
		    disabled = chose_variant)
    with col8:
        try:
            chosen_lvd = float(chosen_data['lvd'])
        except:
            chosen_lvd = 0.0
        lvd = st.number_input('КДР, см',
		    min_value = 0.0, max_value = 15.0,
		    value = chosen_lvd,
		    disabled = chose_variant)
    with col9:
        try:
            chosen_pw = float(chosen_data['pw'])
        except:
            chosen_pw = 0.0
        pw = st.number_input('ЗСЛЖд, см',
		    min_value = 0.0, max_value = 5.0,
		    value = chosen_pw,
		    disabled = chose_variant)
	
	# КСР, КДО, КСО + вычисления объемов методом "пули"
    col10, col11, col12 = st.columns(3)
    with col10:
        try:
            chosen_lvs = float(chosen_data['lvs'])
        except:
            chosen_lvs = 0.0
        lvs = st.number_input('КСР, см',
		    min_value = 0.0, max_value = 15.0,
		    value = chosen_lvs,
		    disabled = chose_variant)
        bullit_lv_dvol = round(7 * lvd**3 / (2.4 + lvd))
        bullit_lv_svol = round(7 * lvs**3 / (2.4 + lvs)) 
    with col11:
        try:
            chosen_lv_dvol = int(chosen_data['lv_dvol'])
        except:
            chosen_lv_dvol = bullit_lv_dvol
        lv_dvol = st.number_input('КДО, мл',
		    value = chosen_lv_dvol,
		    disabled = chose_variant)		    
    with col12:
        try:
            chosen_lv_svol = int(chosen_data['lv_svol'])
        except:
            chosen_lv_svol = bullit_lv_svol
        lv_svol = st.number_input('КСО, мл',
		    value = chosen_lv_svol,
            disabled = chose_variant)
     
     # Вычисление УО, ФВ ЛЖ, ММЛЖ, ИММЛЖ
    col13, col14, col15, col16 = st.columns(4)
    with col13:
        calculate_beat_volume = lv_dvol - lv_svol
        beat_volume = st.number_input('УО, мл',
		    value = calculate_beat_volume,
		    disabled = True)
    with col14:
        try:
            calculate_ef = round((lv_dvol-lv_svol)/lv_dvol * 100)
        except:
            calculate_ef = 0
        ef = st.number_input('ФВ, %',
		    value = calculate_ef,
		    disabled = True)
    with col15:
        try:
            calculate_myocard_mass = round(0.8 * 1.04 * (((lvd + ivs + pw)**3) - lvd**3) + 0.6)
        except:
            calculate_myocard_mass = 0
        myocard_mass = st.number_input('ММЛЖ, г',
		        value = calculate_myocard_mass,
		        disabled = True)
    with col16:
        myocard_mass_index = st.number_input('ИМЛЖ, г/м^2',
		        value = round(calculate_myocard_mass / calculate_body_area),
		        disabled = True)
			
    st.write('PSAX')
    col17, col18, col19 = st.columns(3)
    
    with col17:
        try:
            chosen_pa = float(chosen_data['pa'])
        except:
            chosen_pa = 0.0
        pa = st.number_input('ЛА, см',
		    value = chosen_pa,
            disabled = chose_variant)            
    with col18:
        chosen_pa_valv = str(chosen_data['pa_valv'].tolist())
        chosen_pa_valv = 'Клапан ЛА: ' + chosen_pa_valv[2:-2]
        if chosen_exam == None:
            pa_valv = st.multiselect('Клапан ЛА: ',
	        options = (multiselect['valvs']),
	        default=(multiselect['valvs'].iloc[0]),
		    placeholder='Выберите описание клапана легочной артерии',
		    disabled=chose_variant)
            pa_valv = ', '.join(pa_valv)
        else:			
            pa_valv = st.text_input('Клапан ЛА:', 
			    value = chosen_pa_valv,
			    disabled=chose_variant)
    with col19:
        chosen_pa_valv_kinetic = str(chosen_data['pa_valv_kinetic'].tolist())
        chosen_pa_valv_kinetic = 'Створки КЛА: ' + chosen_pa_valv_kinetic[2:-2]
        if chosen_exam == None:
            pa_valv_kinetic = st.multiselect('Створки КЛА: ',
	        options = (multiselect['valvs_kinetic']),
	        default=(multiselect['valvs_kinetic'].iloc[0]),
		    placeholder='Выберите подвижность створок КЛА',
		    disabled=chose_variant)
            pa_valv_kinetic = ', '.join(pa_valv_kinetic)
        else:			
            pa_valv_kinetic = st.text_input('Створки КЛА:', 
			    value = chosen_pa_valv_kinetic,
			    disabled=chose_variant)		        
    
    chose_pg_vmax = st.selectbox(
	    'Pg max или V max',
		('Pg max', 'V max'), index=0, 
		disabled=chose_variant)
    pg_selected = True	
    if chose_pg_vmax == 'V max':
       pg_selected = False            

    
    # Диаметр легочной артерии, градиент и скорость на клапане легочной артерии
    col20, col21, col22 = st.columns(3)
    
    if pg_selected == True:        
        with col20:        
            try:
                chosen_pg_pa_valv = float(chosen_data['pg_pa_valv'])
            except:
                chosen_pg_pa_valv = 0.0
            pg_pa_valv = st.number_input('Pg max ЛА, mmHg',
		        value = chosen_pg_pa_valv,
                disabled = chose_variant)
        with col21:
            v_max_pa_valv = st.number_input('V max ЛА, м/с',
		        value = round((pg_pa_valv / 4)**0.5, 2),
                disabled = True)                
    elif pg_selected == False:
        with col20:        
            try:
                chosen_v_max_pa_valv = float(chosen_data['v_max_pa_valv'])
            except:
                chosen_v_max_pa_valv = 0.0
            v_max_pa_valv = st.number_input('V max ЛА, м/с',
		        value = chosen_v_max_pa_valv,
                disabled = chose_variant)
        with col21:
            pg_pa_valv = st.number_input('Pg max ЛА, mmHg',
		        value = round(4 * v_max_pa_valv**2, 1),
                disabled = True)
    with col22:
        try:
            chosen_pa_reg = int(chosen_data['pa_reg'])
        except:
            chosen_pa_reg = 1
        pa_reg = st.number_input('Рег. КЛА, ст',
            min_value=0, max_value=4,
		    value = chosen_pa_reg,
            disabled = chose_variant)
    
    # Зоны гипокинеза по сегментам
    st.write('')
    st.image('images/segments.jpg')    
    chosen_kinetic = str(chosen_data['kinetic'].tolist())
    chosen_kinetic = chosen_kinetic[2:-2]        
    if chosen_exam == None:
        kinetic = st.multiselect('Зоны нарушения локальной сократимости миокарда: ',
	    options = (multiselect['kinetic']),
	    default=(multiselect['kinetic'].iloc[0]),
		placeholder='Выберите из списка',
		disabled=chose_variant)
        kinetic = ', '.join(kinetic)
    else:			
        kinetic = st.text_input('Зоны нарушения локальной сократимости миокарда: ',
			chosen_kinetic,
			disabled=chose_variant)
        
    st.write('A4C')
    
    # Левое предсердие апикально, объем ЛП, ИОЛП
    col23, col24, col25, col26 = st.columns(4)
    with col23:
        try:
            chosen_la_x = float(chosen_data['la_x'])
        except:
            chosen_la_x = 0.0
        la_x = st.number_input('ЛПх',
		    value = chosen_la_x,
		    disabled = chose_variant)
    with col24:
        try:
            chosen_la_y = float(chosen_data['la_y'])
        except:
            chosen_la_y = 0.0
        la_y = st.number_input('ЛПy',
		    value = chosen_la_y,
		    disabled = chose_variant)	
    with col25:
        try:
            chosen_la_vol = int(chosen_data['la_vol'])
        except:
            chosen_la_vol = 0
        la_vol = st.number_input('Объем ЛП, мл',
		    value = chosen_la_vol,
		    disabled = chose_variant)
    with col26:
        try:
            la_vol_index = st.number_input('ИОЛП, мл/м^2',
		    value = round(la_vol / calculate_body_area, 1),
		    disabled = True)
        except:
            la_vol_index = st.number_input('ИОЛП, мл/м^2',
		    value = 0.0,
		    disabled = True)	 
	
	# Правое предсердие апикально, объем ЛП, ИОПП
    col27, col28, col29, col30 = st.columns(4)
    with col27:
        try:
            chosen_ra_x = float(chosen_data['ra_x'])
        except:
            chosen_ra_x = 0.0
        ra_x = st.number_input('ППх',
		    value = chosen_ra_x,
		    disabled = chose_variant)
    with col28:
        try:
            chosen_ra_y = float(chosen_data['ra_y'])
        except:
            chosen_ra_y = 0.0
        ra_y = st.number_input('ППy',
		    value = chosen_ra_y,
		    disabled = chose_variant)	
    with col29:
        try:
            chosen_ra_vol = int(chosen_data['ra_vol'])
        except:
            chosen_ra_vol = 0
        ra_vol = st.number_input('Объем ПП, мл',
		    value = chosen_ra_vol,
		    disabled = chose_variant)
    with col30:
        try:
            ra_vol_index = st.number_input('ИОПП, мл/м^2',
		    value = round(ra_vol / calculate_body_area, 1),
		    disabled = True)
        except:
            ra_vol_index = st.number_input('ИОПП, мл/м^2',
		    value = 0.0,
		    disabled = True)	 
    
    # Систолическая функция правого желудочка, диастолическая функция левого
    col31, col32, col33 = st.columns(3)
    with col31:
        try:
            chosen_tapse = int(chosen_data['tapse'])
        except:
            chosen_tapse = 0
        tapse = st.number_input('TAPSE, мм',
		    value = chosen_tapse,
		    disabled = chose_variant)		    
    with col32:
        try:
            chosen_e_peak = int(chosen_data['e_peak'])
        except:
            chosen_e_peak = 0
        e_peak = st.number_input('Пик Е, см/с',
		    value = chosen_e_peak,
		    disabled = dyastola)
    with col33:
        try:
            chosen_a_peak = int(chosen_data['a_peak'])
        except:
            chosen_a_peak = 0
        a_peak = st.number_input('Пик A, см/с',
		    value = chosen_a_peak,
		    disabled = dyastola)
	
    # Диастолическая функция левого желудочка - ТКД
    col34, col35, col36, col37 = st.columns(4)
    with col34:
        try:
            chosen_em_lat = int(chosen_data['em_lat'])
        except:
            chosen_em_lat = 0
        em_lat = st.number_input("e' лат",
		    value = chosen_em_lat,
		    disabled = dyastola)		    
    with col35:
        try:
            chosen_em_med = int(chosen_data['em_med'])
        except:
            chosen_em_med = 0
        em_med = st.number_input("e' мед",
		    value = chosen_em_med,
		    disabled = dyastola)
    with col36:
        try:
            e_a = st.number_input('E/A',
		    value = round(e_peak / a_peak, 1),
		    disabled = True)
        except:
            e_a = st.number_input('E/A',
		    value = 0.0,
		    disabled = True)
    with col37:
        try:
            e_em = st.number_input("E/e' мед",
		    value = round(e_peak / em_med, 1),
		    disabled = True)
        except:
            e_em = st.number_input("E/e' мед",
		    value = 0.0,
		    disabled = True)
	
	# Клапанный аппарат
	# Аортальный клапан	 	
    col38, col39 = st.columns(2)
    with col38:
        chosen_ao_valv = str(chosen_data['ao_valv'].tolist())
        chosen_ao_valv = chosen_ao_valv[2:-2]
        if chosen_exam == None:
            ao_valv = st.multiselect('Аортальный клапан: ',
	        options = (multiselect['valvs']),
	        default=(multiselect['valvs'].iloc[1]),
		    disabled=chose_variant)
            ao_valv = ', '.join(ao_valv)
        else:			
            ao_valv = st.text_input('Аортальный клапан:', 
			    value = chosen_ao_valv,
			    disabled=chose_variant)
    with col39:
        chosen_ao_valv_kinetic = str(chosen_data['ao_valv_kinetic'].tolist())
        chosen_ao_valv_kinetic = chosen_ao_valv_kinetic[2:-2]
        if chosen_exam == None:
            ao_valv_kinetic = st.multiselect('Створки АК: ',
	        options = (multiselect['valvs_kinetic']),
	        default=(multiselect['valvs_kinetic'].iloc[0]),
		    disabled=chose_variant)
            ao_valv_kinetic = ', '.join(ao_valv_kinetic)
        else:			
            ao_valv_kinetic = st.text_input('Створки АК:', 
			    value = chosen_ao_valv_kinetic,
			    disabled=chose_variant)		        
		
    col40, col41, col42 = st.columns(3)   
    if pg_selected == True:        
        with col40:        
            try:
                chosen_pg_ao_valv = float(chosen_data['pg_ao_valv'])
            except:
                chosen_pg_ao_valv = 0.0
            pg_ao_valv = st.number_input('Pg max АК, mmHg',
		        value = chosen_pg_ao_valv,
                disabled = chose_variant)
        with col41:
            v_max_ao_valv = st.number_input('V max АК, м/с',
		        value = round((pg_ao_valv / 4)**0.5, 1),
                disabled = True)                
    elif pg_selected == False:
        with col40:        
            try:
                chosen_v_max_ao_valv = float(chosen_data['v_max_ao_valv'])
            except:
                chosen_v_max_ao_valv = 0.0
            v_max_ao_valv = st.number_input('V max АК, м/с',
		        value = chosen_v_max_ao_valv,
                disabled = chose_variant)
        with col41:
            pg_ao_valv = st.number_input('Pg max АК, mmHg',
		        value = round(4 * v_max_ao_valv**2, 1),
                disabled = True)    
    with col42:
        try:
            chosen_ar = int(chosen_data['ar'])
        except:
            chosen_ar = 0
        ar = st.number_input('Рег. АК, ст',
		    value = chosen_ar,
		    min_value=0, max_value=4,
            disabled = chose_variant)

	# Митральный клапан	    
    col43, col44 = st.columns(2)
    with col43:
        chosen_m_valv = str(chosen_data['m_valv'].tolist())
        chosen_m_valv = chosen_m_valv[2:-2]
        if chosen_exam == None:
            m_valv = st.multiselect('Митральный клапан: ',
	        options = (multiselect['valvs']),
	        default=(multiselect['valvs'].iloc[1]),
		    disabled=chose_variant)
            m_valv = ', '.join(m_valv)
        else:			
            m_valv = st.text_input('Митральный клапан:', 
			    value = chosen_m_valv,
			    disabled=chose_variant)
    with col44:
        chosen_m_valv_kinetic = str(chosen_data['m_valv_kinetic'].tolist())
        chosen_m_valv_kinetic = chosen_m_valv_kinetic[2:-2]
        if chosen_exam == None:
            m_valv_kinetic = st.multiselect('Створки МК: ',
	        options = (multiselect['valvs_kinetic']),
	        default=(multiselect['valvs_kinetic'].iloc[0]),
		    placeholder='Выберите подвижность створок МК',
		    disabled=chose_variant)
            m_valv_kinetic = ', '.join(m_valv_kinetic)
        else:			
            m_valv_kinetic = st.text_input('Створки МК:', 
			    value = chosen_m_valv_kinetic,
			    disabled=chose_variant)		        
		
    col45, col46, col47 = st.columns(3)   
    if pg_selected == True:        
        with col45:        
            try:
                chosen_pg_m_valv = float(chosen_data['pg_m_valv'])
            except:
                chosen_pg_m_valv = 0.0
            pg_m_valv = st.number_input('Pg max МК, mmHg',
		        value = chosen_pg_m_valv,
                disabled = chose_variant)
        with col46:
            v_max_m_valv = st.number_input('V max МК, м/с',
		        value = round((pg_m_valv / 4)**0.5, 1),
                disabled = True)                
    elif pg_selected == False:
        with col45:        
            try:
                chosen_v_max_m_valv = float(chosen_data['v_max_m_valv'])
            except:
                chosen_v_max_m_valv = 0.0
            v_max_m_valv = st.number_input('V max МК, м/с',
		        value = chosen_v_max_m_valv,
                disabled = chose_variant)
        with col46:
            pg_m_valv = st.number_input('Pg max МК, mmHg',
		        value = round(4 * v_max_m_valv**2, 1),
                disabled = True)    
    with col47:
        try:
            chosen_mr = int(chosen_data['mr'])
        except:
            chosen_mr = 1
        mr = st.number_input('Рег. МК, ст',
            min_value=0, max_value=4,
		    value = chosen_mr,
            disabled = chose_variant)
    
	# Трикуспидальный клапан
    col48, col49 = st.columns(2)
    with col48:
        chosen_t_valv = str(chosen_data['t_valv'].tolist())
        chosen_t_valv = chosen_t_valv[2:-2]
        if chosen_exam == None:
            t_valv = st.multiselect('Трикуспидальный клапан: ',
	        options = (multiselect['valvs']),
	        default=(multiselect['valvs'].iloc[0]),
		    disabled=chose_variant)
            t_valv = ', '.join(t_valv)
        else:			
            t_valv = st.text_input('Трикуспидальный клапан:', 
			    value = chosen_t_valv,
			    disabled=chose_variant)
    with col49:
        chosen_t_valv_kinetic = str(chosen_data['t_valv_kinetic'].tolist())
        chosen_t_valv_kinetic = chosen_t_valv_kinetic[2:-2]
        if chosen_exam == None:
            t_valv_kinetic = st.multiselect('Створки ТК: ',
	        options = (multiselect['valvs_kinetic']),
	        default=(multiselect['valvs_kinetic'].iloc[0]),
		    disabled=chose_variant)
            t_valv_kinetic = ', '.join(t_valv_kinetic)
        else:			
            t_valv_kinetic = st.text_input('Створки ТК:', 
			    value = chosen_t_valv_kinetic,
			    disabled=chose_variant)		        
		
    col50, col51, col52 = st.columns(3)   
    if pg_selected == True:        
        with col50:        
            try:
                chosen_pg_t_valv = int(chosen_data['pg_t_valv'])
            except:
                chosen_pg_t_valv = 0
            pg_t_valv = st.number_input('Pg max ТР, mmHg',
		        value = chosen_pg_t_valv,
                disabled = chose_variant)
        with col51:
            v_max_t_valv = st.number_input('V max ТР, м/с',
		        value = round((pg_t_valv / 4)**0.5, 1),
                disabled = True)                
    elif pg_selected == False:
        with col50:        
            try:
                chosen_v_max_t_valv = float(chosen_data['v_max_t_valv'])
            except:
                chosen_v_max_t_valv = 0.0
            v_max_t_valv = st.number_input('V max ТК, м/с',
		        value = chosen_v_max_t_valv,
                disabled = chose_variant)
        with col51:
            pg_t_valv = st.number_input('Pg max ТК, mmHg',
		        value = round(4 * v_max_t_valv**2, 1),
                disabled = True)    
    with col52:
        try:
            chosen_tr = int(chosen_data['tr'])
        except:
            chosen_tr = 1
        tr = st.number_input('Рег. ТК, ст',
		    value = chosen_tr,
		    min_value=0, max_value=4,
            disabled = chose_variant)
	
	# Прочее	
    chosen_others = str(chosen_data['others'].tolist())
    if chosen_others == '[nan]':
        chosen_others = ''
    else:    
        chosen_others = chosen_others[2:-2]	
    if chose_variant == False or check_redact_exam == True:
        others = st.text_area('Прочее', 
			    value = '',
			    disabled=chose_variant)		      
    else:			
        others = st.text_area('Прочее', 
			value = chosen_others,
			disabled=chose_variant)		      
	
	# Признаки гидроперикарда		        			    		
    chosen_pericard = str(chosen_data['pericard'].tolist())
    chosen_pericard = chosen_pericard[2:-2]
    if chose_variant == False or check_redact_exam == True:
        pericard = st.text_input('Признаки гидроперикарда: ', 
			    value = chosen_pericard,
			    disabled=chose_variant)		      
    else:			
        pericard = st.text_input('Признаки гидроперикарда: ', 
			value = chosen_pericard,
			disabled=chose_variant)
    if pericard == '':
        pericard = 'нет. '
			
	# Нижняя полая вена
    st.write('SUB')
    col53, col54 = st.columns(2)
    with col53:
        try:
            chosen_cv = float(chosen_data['cv'])
        except:
            chosen_cv = 0.0
        cv = st.number_input('Диаметр НПВ, см',
		        value = chosen_cv,
		        disabled = chose_variant)
    with col54:
        chosen_cv_colab = str(chosen_data['cv_colab'].tolist())
        chosen_cv_colab = 'Коллабирование НПВ на вдохе: ' + chosen_cv_colab[2:-2]
        if chose_variant == False or check_redact_exam == True:
            cv_colab = st.selectbox('Выберите степень коллабирования НПВ на вдохе',
		           ('более 50%', 'менее 50%', 'не коллабирует'), index=0, 
				   disabled=chose_variant)
        else:			
            cv_colab = st.text_input('Коллабирование НПВ на вдохе: ', 
			    value = chosen_cv_colab,
			    disabled=chose_variant)	
		
# ПРОГНОЗ НЕЙРОСЕТИ_____________________________________________________
    
    st.write('')
    st.write('ПРОГНОЗ НЕЙРОСЕТИ')

    col55, col56 = st.columns(2)
    with col55:
            if st.button("Получить прогноз нейросети AF-Predict 2.0"):
                try:
                    result = prediction(ao, la, lvd, rv, ra_x, pa, ivs, pw, ef, ar, mr, tr, cv)
                    out = round(result[0][0], 2)
                except:
                    st.error('Проверьте правильность заполнения формы!', icon="⚠️")
    with col56:
        try:		
            st.success(str(out), icon="✅")
        except:
            pass

    try:
        result = prediction(ao, la, lvd, rv, ra_x, pa, ivs, pw, ef, ar, mr, tr, cv)
        nw_prognose = round(result[0][0], 2)
    except:
        nw_prognose = None				

# ФОРМИРОВАНИЕ ПРОТОКОЛА ИССЛЕДОВАНИЯ

	# Автоматическое заключение
	# Состояние стенки аорты, аортального и митрального клапанов
    if aorta_wall != 'не изменена':
        aorta_wall_report = 'Уплотнение стеки аорты. '
    else:
        aorta_wall_report = ''
    if ao_valv != 'не изменены':
        ao_valv_report = 'Уплотнение створок аортального клапана. '
    else:
        ao_valv_report = ''
    if ao_valv_kinetic != 'подвижность не ограничена':
        ao_valv_kinetic_report = 'Ограничение подвижности створок аортального клапана. '
    else:
	    ao_valv_kinetic_report = ''
    if m_valv != 'не изменены':
        m_valv_report = 'Уплотнение створок митрального клапана. '
    else:
        m_valv_report = ''
    if m_valv_kinetic != 'подвижность не ограничена':
        m_valv_kinetic_report = 'Ограничение подвижности створок митрального клапана. '
    else:
	    m_valv_kinetic_report = ''
    if t_valv != 'не изменены':
        t_valv_report = 'Уплотнение створок трикуспидального клапана. '
    else:
        t_valv_report = ''
    if t_valv_kinetic != 'подвижность не ограничена':
        t_valv_kinetic_report = 'Ограничение подвижности створок трикуспидального клапана. '
    else:
	    t_valv_kinetic_report = ''
    if pa_valv != 'не изменены':
        pa_valv_report = 'Уплотнение створок клапана легочной артерии. '
    else:
        t_valv_report = ''
    if t_valv_kinetic != 'подвижность не ограничена':
        t_valv_kinetic_report = 'Ограничение подвижности створок клапана легочной артерии. '
    else:
	    t_valv_kinetic_report = ''
	# Размеры камер сердца
    if ao_valsalva > 4.0 and ao > 4.0:
        ao_report = 'Дилатация корня и восходящего отдела аорты. '
    elif ao_valsalva > 4.0 and ao <= 4.0:
        ao_report = 'Дилатация корня аорты. '
    elif ao_valsalva <= 4.0 and ao > 4.0:
	    ao_report = 'Дилатация восходящего отдела аорты. '	  
    else:
        ao_report = ''
    if (la > 4.0 or la_x > 4.5 or la_y > 5.3 or la_vol > 52) and la_vol_index > 28:
        la_dyl = 1
    else:
        la_dyl = 0
    if lvd > 5.6:
        lv_dyl = 1
    else:
        lv_dyl = 0
    if ra_x > 3.9 or ra_y > 4.6:
        ra_dyl = 1
    else:
        ra_dyl = 0
    if rv > 2.9:
        rv_dyl = 1
    else:
        rv_dyl = 0
    if pa > 2.9:
        pa_report = 'Дилатация ствола легочной артерии. '
    else:
        pa_report = ''
    if cv > 2.6:
        cv_report = 'Дилатация нижней полой вены. '
    else:
        cv_report = ''
    if la_dyl == 1 and lv_dyl == 1 and ra_dyl == 1 and rv_dyl == 1:
        chambers_report = 'Дилатация всех камер сердца. '
    elif la_dyl == 1 and lv_dyl == 1 and ra_dyl == 1 and rv_dyl == 0:
        chambers_report = 'Дилатация левых камер сердца и правого предсердия. '
    elif la_dyl == 1 and lv_dyl == 1 and ra_dyl == 0 and rv_dyl == 1:
        chambers_report = 'Дилатация левых камер сердца и правого желудочка. '
    elif la_dyl == 1 and lv_dyl == 0 and ra_dyl == 1 and rv_dyl == 1:
        chambers_report = 'Дилатация правых камер сердца и левого предсердия. '
    elif la_dyl == 0 and lv_dyl == 1 and ra_dyl == 1 and rv_dyl == 1:
        chambers_report = 'Дилатация правых камер сердца и левого желудочка. '
    elif la_dyl == 1 and lv_dyl == 0 and ra_dyl == 1 and rv_dyl == 0:
        chambers_report = 'Дилатация обоих предсердий. '
    elif la_dyl == 0 and lv_dyl == 1 and ra_dyl == 0 and rv_dyl == 1:
        chambers_report = 'Дилатация обоих желудочков. '
    elif la_dyl == 1 and lv_dyl == 1 and ra_dyl == 0 and rv_dyl == 0:
        chambers_report = 'Дилатация левых камер сердца. '
    elif la_dyl == 0 and lv_dyl == 0 and ra_dyl == 1 and rv_dyl == 1:
        chambers_report = 'Дилатация правых камер сердца. '
    elif la_dyl == 1 and lv_dyl == 0 and ra_dyl == 0 and rv_dyl == 1:
        chambers_report = 'Дилатация левого предсердия и правого желудочка. '
    elif la_dyl == 0 and lv_dyl == 1 and ra_dyl == 1 and rv_dyl == 0:
        chambers_report = 'Дилатация правого предсердия и левого желудочка. '        
    elif la_dyl == 1 and lv_dyl == 0 and ra_dyl == 0 and rv_dyl == 0:
        chambers_report = 'Дилатация левого предсердия. '
    elif la_dyl == 0 and lv_dyl == 1 and ra_dyl == 0 and rv_dyl == 0:
        chambers_report = 'Дилатация левого желудочка. '
    elif la_dyl == 0 and lv_dyl == 0 and ra_dyl == 1 and rv_dyl == 0:
        chambers_report = 'Дилатация правого предсердия. '
    elif la_dyl == 0 and lv_dyl == 0 and ra_dyl == 0 and rv_dyl == 1:
        chambers_report = 'Дилатация правого желудочка. '
    else:
        chambers_report = 'Размеры камер сердца в норме. '
    # Толщина стенок сердца
    if ivs > 1.1 and pw > 1.1:
        septs_report = 'Концентрическая гипертрофия миокарда левого желудочка. '
    elif ivs > 1.1 or pw > 1.1:
        septs_report = 'Эксцентрическая гипертрофия миокарда левого желудочка. '
    else:
        septs_report = ''
    if rv_wall > 0.4:
        rv_wall_report = 'Гипертрофия миокарда правого желудочка. '
    else:
        rv_wall_report = ''
    if septs_report == '' and rv_wall_report == '':
        septs_report = 'Толщина стенок сердца в пределах нормы. '
    # Кинетика миокарда, систолическая и диастолическая функция ЛЖ
    if ef >= 50:
        ef_report = 'Систолическая функция левого желудочка не нарушена. '
    elif ef < 50 and ef >= 40:
        ef_report = 'Систолическая функция левого желудочка незначительно снижена. '
    elif ef <40 and ef >= 30:
        ef_report = 'Систолическая функция левого желудочка умеренно снижена. '
    else:
        ef_report = 'Систолическая функция левого желудочка значительно снижена. '
        
    if rythm != 'Ритм синусовый':
        df_report = 'Диастолическая функция левого желудочка не определена. '
    elif em_lat >= 10 or em_med >= 8:
        df_report = 'Диастолическая функция левого желудочка не нарушена. '
    elif (em_lat <10 or em_med <8) and e_a <= 0.8:
        df_report = 'Диастолическая функция левого желудочка нарушена по 1 типу (нарушение релаксации). '
    elif (em_lat <10 or em_med <8) and e_a > 2.0:
        df_report = 'Диастолическая функция левого желудочка нарушена по 2 типу (рестриктивному). '
    else:
        df_report = 'Диастолическая функция левого желудочка нарушена по типу псевдонормализации. '
   
    if tapse >= 15:
        rv_function_report = 'Систолическая функция правого желудочка сохранена. '
    else:
        rv_function_report = 'Систолическая функция правого желудочка снижена. '
    # Клапанный аппарат
    if pg_ao_valv > 60 or v_max_ao_valv > 3.9:
        pg_ao_valv_report = 'УЗ-признаки тяжелого аортального стеноза. '
    elif (pg_ao_valv <= 60 and pg_ao_valv > 40) or (v_max_ao_valv <= 3.9 and v_max_ao_valv > 3.2):
        pg_ao_valv_report = 'УЗ-признаки умеренного аортального стеноза. '
    elif (pg_ao_valv <= 40 and pg_ao_valv > 20) or (v_max_ao_valv <= 3.2 and v_max_ao_valv > 2.2):
        pg_ao_valv_report = 'УЗ-признаки незначительного аортального стеноза. '
    else:
        pg_ao_valv_report = ''
    if ar == 4:
        ar_report = 'Тяжелая аортальная регургитация. '
    elif ar == 3:
        ar_report = 'Значительная аортальная регургитация. '
    elif ar == 2:
        ar_report = 'Умеренная аортальная регургитация. '
    elif ar == 1:
        ar_report = 'Незначительная аортальная регургитация. '
    else:
        ar_report = ''                
    if pg_m_valv > 20 or v_max_ao_valv > 2.2:
        pg_m_valv_report = 'УЗ-признаки тяжелого митрального стеноза. '
    elif (pg_m_valv <= 20 and pg_m_valv > 12) or (v_max_m_valv <= 2.2 and v_max_m_valv > 1.7):
        pg_m_valv_report = 'УЗ-признаки умеренного митрального стеноза. '
    elif (pg_m_valv <= 12 and pg_m_valv > 7) or (v_max_m_valv <= 1.7 and v_max_m_valv > 1.3):
        pg_m_valv_report = 'УЗ-признаки незначительного митрального стеноза. '
    else:
        pg_m_valv_report = ''
    if mr == 4:
        mr_report = 'Тяжелая митральная регургитация. '
    elif mr == 3:
        mr_report = 'Значительная митральная регургитация. '
    elif mr == 2:
        mr_report = 'Умеренная митральная регургитация. '
    elif mr == 1:
        mr_report = 'Незначительная митральная регургитация. '
    else:
        mr_report = ''        
    if tr == 4:
        tr_report = 'Тяжелая трикуспидальная регургитация. '
    elif tr == 3:
        tr_report = 'Значительная трикуспидальная регургитация. '
    elif tr == 2:
        tr_report = 'Умеренная трикуспидальная регургитация. '
    elif tr == 1:
        tr_report = 'Незначительная трикуспидальная регургитация. '
    else:
        tr_report = ''        
    if pg_pa_valv > 10 or v_max_pa_valv > 1.6:
        pg_pa_valv_report = 'УЗ-признаки стеноза клапана легочной артерии. '
    else:
        pg_pa_valv_report = ''
    if pa_reg == 4:
        pa_reg_report = 'Тяжелая регургитация на клапане легочной артерии. '
    elif pa_reg == 3:
        pa_reg_report = 'Значительная регургитация на клапане легочной артерии. '
    elif pa_reg == 2:
        pa_reg_report = 'Умеренная регургитация на клапане легочной артерии. '
    elif pa_reg == 1:
        pa_reg_report = 'Незначительная регургитация на клапане легочной артерии. '
    else:
        pa_reg_report = ''                	
	# Расчетное ДЛА
    if cv_colab == 'более 50%' and cv <= 2.3:
        pa_press = pg_t_valv
    elif cv_colab != 'более 50%' and cv <= 2.3:
        pa_press = pg_t_valv + 5
    elif cv_colab == 'более 50%' and cv > 2.3:
        pa_press = pg_t_valv + 10
    elif cv_colab != 'более 50%' and cv > 2.3:
        pa_press = pg_t_valv + 15
    else:
        pa_press = 0
    try:		
        pa_pressure = 'Расчетное давление в легочной артерии ' + str(pa_press) + '-' + str(pa_press + 5) + ' mmHg. '
    except:
        pa_pressure	= 'Давление в легочной артерии не определяется. '
    if pa_press < 35:
        pa_press_report = ''
    elif pa_press >= 35 and pa_press < 50:
        pa_press_report = 'Признаки начальной легочной гипертензии. '
    elif pa_press >= 50 and pa_press < 80:
        pa_press_report = 'Признаки умеренной легочной гипертензии. '
    elif pa_press >= 80:
        pa_press_report = 'Признаки значительной легочной гипертензии. '
    else:
        pa_press_report	= ''			    
    # Перикард и прочее
    if pericard == 'нет. ' or pericard == '':
        pericard_report = ''
    else:
        pericard_report = 'УЗ-признаки наличия жидкости в полости перикарда. '
    if others == '':
        others_report = ''
    else:
        others_report = '\n\n' + others
			
	# АВТООТЧЕТ
    id_of_exam = str(name_of_patient) + ' ' + str(date_of_birth) + ' г.р., ТТЭ от ' + str(date_of_exam) + ' ' + str(time_of_exam)
    header = 'ГУЗ “Липецкая городская больница №4 “Липецк-Мед”, отделение функциональной диагностики'
    patient = 'Ф.И.О.: ' + str(name_of_patient) + ' ' + str(date_of_birth)
    exam = 'Дата и время исследования: ' + str(date_of_exam) + ' ' + str(time_of_exam)
    tte_scanner = scanner
    troubles = 'Ограничение визуализации: ' + visualise_troubles
    ryth_height_weight = rythm + '. Рост - ' + str(patient_height) + ' см. Вес - ' + str(patient_weight) + ' кг.'
    tte_number = '\n' + 'ТРАНСТОРАКАЛЬНАЯ ЭХОКАРДИОГРАФИЯ № ' + str(exam_number) + '\n'
    protocol_data_a = id_of_exam + '\n' + header + '\n' + patient + '\n' + exam + '\n' + tte_scanner + '\n' + troubles + '\n' + ryth_height_weight + '\n' + tte_number + '\n'

    aorta = 'Стенка аорты ' + aorta_wall + '. Аорта на уровне синуса Вальсальвы  ' + str(ao_valsalva) + ' см (норме менее 4.0 см). Восходящая аорта ' + str(ao) + ' см (норме менее 4.0 см).'
    left_atrium = 'Левое предсердие PLAX ' + str(la) + ' см (норма менее 4.0 см). Левое предсердие апикально ' + str(la_x) + ' х ' + str(la_y) + ' см (норма менее 4.5 х 5.3 см). Объем левого предсердия ' + str(la_vol) + ' мл (норма менее 52 мл). Индекс объема ЛП ' + str(la_vol_index) + ' мл/м^2 (норма 16 - 28 мл/м^2).'
    left_ventricle = 'КДР ЛЖ ' + str(lvd) + ' см (норма менее 5.6 см). КСР ЛЖ ' + str(lvs) + ' см (норма менее 2.8 см). КДО ЛЖ ' + str(lv_dvol) + ' мл (норма менее 104 мл). КСО ЛЖ ' + str(lv_svol) + ' мл (норма менее 49 мл).'
    walls = 'Толщина МЖП ' + str(ivs) + ' см (норма менее 1.1 см). Толщина ЗСЛЖ ' + str(pw) + ' см (норма менее 1.1 см). ММЛЖ ' + str(myocard_mass) + ' г (норма -   женщины менее 162 г, мужчины менее 224 г). Индекс ММЛЖ ' + str(myocard_mass_index) + ' кг/м^2 (норма 44 - 88 г/м2).'  
    right_atrium = 'Правое предсердие A4C ' + str(ra_x) + ' x ' + str(ra_y) + ' см (норма менее 3.9 х 4.6 см). Правый желудочек ' + str(rv) + ' см (норма менее 3.0 см). Толщина стенки ПЖ ' + str(rv_wall) + ' см (норма менее 0.4 см). Легочная артерия ' + str(pa) + ' см. (норма менее 2.9 см) Нижняя полая вена ' + str(cv) + ' см (норма менее 2.6 см). НПВ на вдохе коллабирует ' + str(cv_colab) + '.'
    if rythm == 'Ритм синусовый':
        functions = str(kinetic) + ' ФВ ЛЖ (Симпсон) ' + str(ef) + '% (норма более 50%). TAPSE ПЖ ' + str(tapse) + ' мм (норма более 15 мм). Трансмитральный поток: пик Е ' + str(e_peak) + ' см/с (норма менее 120 см/с), пик А ' + str(a_peak) + ' см/с (норма менее 120 см/с), Пик E / Пик A ' + str(e_a) + " (норма 0.8 - 2.0). ТИД: e' лат " + str(em_lat) + " см/с (норма более 10 см/с).  e' мед " + str(em_med) + " см/с (норма более 8 см/с). Е/е' " + str(e_em) + " (норма менее 8). " 
    else:
        functions = str(kinetic) + ' ФВ ЛЖ (Симпсон) ' + str(ef) + '% (норма более 50%). TAPSE ПЖ ' + str(tapse) + ' мм (норма более 15 мм). ' 
    a_valv = 'Аортальный клапан. Створки ' + str(ao_valv) + '. Подвижность створок ' + str(ao_valv_kinetic) + '. Раскрытие створок АК ' + str(av_plax) + ' см (норма более 1.6 см). Pg max АК ' + str(pg_ao_valv) + ' mmHg (норма менее 10 mmHg). Vmax АК ' + str(v_max_ao_valv) + ' м/с (норма менее 1,6 м/с). Аортальная регургитация ' + str(ar) + ' ст. '
    m_valv = 'Митральный клапан. Створки ' + str(m_valv) +  '. Подвижность створок ' + str(m_valv_kinetic) + '. Pg max МК ' + str(pg_m_valv) + ' mmHg (норма менее 7 mmHg). Vmax МК ' + str(v_max_m_valv) + ' м/с (норма менее 1.3 м/с). Митральная регургитация ' + str(mr) + ' ст. '
    t_valv = 'Трикуспидальный клапан. Створки ' + str(t_valv) +  '. Подвижность створок ' + str(t_valv_kinetic) + '. Трикуспидальная регургитация ' + str(tr) + ' ст. Pg max ТР ' + str(pg_t_valv) + ' mmHg (норма менее 35 mmHg). Vmax ТР ' + str(v_max_t_valv) + ' м/с (норма менее 2.9 м/с). ' + pa_pressure
    p_valv = 'Клапан легочной артерии. Створки ' + str(pa_valv) +  '. Подвижность створок ' + str(pa_valv_kinetic) + '. Pg max КЛА ' + str(pg_pa_valv) + ' mmHg (норма менее 10 mmHg). Vmax КЛА ' + str(v_max_pa_valv) + ' м/с (норма менее 1.6 м/с). Регургитация КЛА ' + str(pa_reg) + ' ст. '
    pericardium = 'Признаки наличия жидкости в полости перикарда: ' + pericard
    protocol_data_b = aorta + '\n' + left_atrium + '\n' + left_ventricle + '\n' + walls + '\n' + right_atrium + '\n' + functions + '\n' + a_valv + '\n' + m_valv + '\n' + t_valv + '\n' + p_valv + '\n' + pericardium + '\n' + others + '\n'
    
    walls_status = aorta_wall_report + ao_valv_report + ao_valv_kinetic_report + m_valv_report + m_valv_kinetic_report + t_valv_report + t_valv_kinetic_report
    chambers_status = ao_report + chambers_report + pa_report + cv_report + septs_report + rv_wall_report	
    function_status = str(kinetic) + ' ' + ef_report + df_report + rv_function_report
    valvs_status = pg_ao_valv_report + ar_report + pg_m_valv_report + mr_report + tr_report + pg_pa_valv_report + pa_reg_report + pa_press_report
    pericard_and_others_status = pericard_report + others_report
    
    protocol_data_c = '\nЗАКЛЮЧЕНИЕ. ' + '\n' + walls_status + chambers_status + function_status + valvs_status + pericard_and_others_status
    
    protocol_data = protocol_data_a + protocol_data_b + protocol_data_c + '\n\nДанное заключение не является диагнозом, необходимо сопоставление с клиническими и лабораторными данными\n\nВрач: Любавин А. В.'
        
    st.write('')
    st.write('ПРОТОКОЛ ИССЛЕДОВАНИЯ')
    chosen_protocol = str(chosen_data['protocol'].tolist())
    chosen_protocol = chosen_protocol[2:-2]
    if chosen_exam == None or check_redact_exam == True:
        protocol = st.text_area('Протокол исследования', 
                            value=protocol_data, 
							height=1000)
    else:
        protocol = st.write(chosen_protocol)
    # СОХРАНЕНИЕ ИЗМЕНЕНИЙ
    
    # Формирование названия файла
    report_filename = 'documents/' + id_of_exam.replace(':', '-') + '.docx'
#    st.write(report_filename)

    if chosen_exam != None:
        col57, col58 = st.columns(2)
        with col57:
            if st.button('Сохранить изменения в исследовании',
						disabled=chose_variant):
                    try:
                        save_exam_list  = [id_of_exam, name_of_patient, date_of_birth, date_of_exam, time_of_exam, patient_height, patient_weight, scanner, rythm, visualise_troubles, aorta_wall, ao_valsalva, ao, av_plax, la, rv, rv_wall, ivs, lvd, pw, lvs, lv_dvol, lv_svol, beat_volume,	ef,	myocard_mass, myocard_mass_index, pa,	pa_valv, pa_valv_kinetic, pg_pa_valv, v_max_pa_valv, pa_reg, kinetic, ra_x,	ra_y, ra_vol, la_x,	la_y, la_vol, tapse, e_peak, a_peak, em_lat, em_med, ao_valv, ao_valv_kinetic, pg_ao_valv,	v_max_ao_valv, ar, m_valv, m_valv_kinetic, pg_m_valv, v_max_m_valv,	mr,	t_valv,	t_valv_kinetic,	pg_t_valv, v_max_t_valv, tr, others, pericard, cv, cv_colab, pa_press, nw_prognose, protocol]
                        data.loc[chosen_exam] = [id_of_exam, name_of_patient, date_of_birth, date_of_exam, time_of_exam, patient_height, patient_weight, scanner, rythm, visualise_troubles, aorta_wall, ao_valsalva, ao, av_plax, la, rv, rv_wall, ivs, lvd, pw, lvs, lv_dvol, lv_svol, beat_volume,	ef,	myocard_mass, myocard_mass_index, pa,	pa_valv, pa_valv_kinetic, pg_pa_valv, v_max_pa_valv, pa_reg, kinetic, ra_x,	ra_y, ra_vol, la_x,	la_y, la_vol, tapse, e_peak, a_peak, em_lat, em_med, ao_valv, ao_valv_kinetic, pg_ao_valv,	v_max_ao_valv, ar, m_valv, m_valv_kinetic, pg_m_valv, v_max_m_valv,	mr,	t_valv,	t_valv_kinetic,	pg_t_valv, v_max_t_valv, tr, others, pericard, cv, cv_colab, pa_press, nw_prognose, protocol]
                        data.drop(exam_number-1, inplace=True)
                        data.to_excel('tabs/exams.xlsx', index=False)
                        report = docx.Document()
                        report.add_heading('', level=1)
                        report.add_paragraph(protocol_data)
                        report.save(report_filename)
                        st.success('Изменения в исследовании сохранены', icon="✅")
                    except:
                        st.error('Сохранение не произведено!', icon="⚠️")       
        with col58:                
            if st.button('Сохранить как новое исследование', 
						disabled=chose_variant):
                    try:
                        save_exam_list  = [id_of_exam, name_of_patient, date_of_birth, date_of_exam, time_of_exam, patient_height, patient_weight, scanner, rythm, visualise_troubles, aorta_wall, ao_valsalva, ao, av_plax, la, rv, rv_wall, ivs, lvd, pw, lvs, lv_dvol, lv_svol, beat_volume,	ef,	myocard_mass, myocard_mass_index, pa,	pa_valv, pa_valv_kinetic, pg_pa_valv, v_max_pa_valv, pa_reg, kinetic, ra_x,	ra_y, ra_vol, la_x,	la_y, la_vol, tapse, e_peak, a_peak, em_lat, em_med, ao_valv, ao_valv_kinetic, pg_ao_valv,	v_max_ao_valv, ar, m_valv, m_valv_kinetic, pg_m_valv, v_max_m_valv,	mr,	t_valv,	t_valv_kinetic,	pg_t_valv, v_max_t_valv, tr, others, pericard, cv, cv_colab, pa_press, nw_prognose, protocol]
                        data.loc[int(len(data)) + 1] = [id_of_exam, name_of_patient, date_of_birth, date_of_exam, time_of_exam, patient_height, patient_weight, scanner, rythm, visualise_troubles, aorta_wall, ao_valsalva, ao, av_plax, la, rv, rv_wall, ivs, lvd, pw, lvs, lv_dvol, lv_svol, beat_volume,	ef,	myocard_mass, myocard_mass_index, pa,	pa_valv, pa_valv_kinetic, pg_pa_valv, v_max_pa_valv, pa_reg, kinetic, ra_x,	ra_y, ra_vol, la_x,	la_y, la_vol, tapse, e_peak, a_peak, em_lat, em_med, ao_valv, ao_valv_kinetic, pg_ao_valv,	v_max_ao_valv, ar, m_valv, m_valv_kinetic, pg_m_valv, v_max_m_valv,	mr,	t_valv,	t_valv_kinetic,	pg_t_valv, v_max_t_valv, tr, others, pericard, cv, cv_colab, pa_press, nw_prognose, protocol]
                        data.to_excel('tabs/exams.xlsx', index=False)
                        report = docx.Document()
                        report.add_heading('', level=1)
                        report.add_paragraph(protocol_data)
                        report.save(report_filename)
                        st.success('Изменения в исследовании сохранены', icon="✅")
                    except:
                        st.error('Сохранение не произведено!', icon="⚠️")
    else:
        if st.button('Сохранить изменения в исследовании'):
                    try:
                        save_exam_list  = [id_of_exam, name_of_patient, date_of_birth, date_of_exam, time_of_exam, patient_height, patient_weight, scanner, rythm, visualise_troubles, aorta_wall, ao_valsalva, ao, av_plax, la, rv, rv_wall, ivs, lvd, pw, lvs, lv_dvol, lv_svol, beat_volume,	ef,	myocard_mass, myocard_mass_index, pa,	pa_valv, pa_valv_kinetic, pg_pa_valv, v_max_pa_valv, pa_reg, kinetic, ra_x,	ra_y, ra_vol, la_x,	la_y, la_vol, tapse, e_peak, a_peak, em_lat, em_med, ao_valv, ao_valv_kinetic, pg_ao_valv,	v_max_ao_valv, ar, m_valv, m_valv_kinetic, pg_m_valv, v_max_m_valv,	mr,	t_valv,	t_valv_kinetic,	pg_t_valv, v_max_t_valv, tr, others, pericard, cv, cv_colab, pa_press, nw_prognose, protocol]
                        data.loc[exam_number - 1] = [id_of_exam, name_of_patient, date_of_birth, date_of_exam, time_of_exam, patient_height, patient_weight, scanner, rythm, visualise_troubles, aorta_wall, ao_valsalva, ao, av_plax, la, rv, rv_wall, ivs, lvd, pw, lvs, lv_dvol, lv_svol, beat_volume,	ef,	myocard_mass, myocard_mass_index, pa,	pa_valv, pa_valv_kinetic, pg_pa_valv, v_max_pa_valv, pa_reg, kinetic, ra_x,	ra_y, ra_vol, la_x,	la_y, la_vol, tapse, e_peak, a_peak, em_lat, em_med, ao_valv, ao_valv_kinetic, pg_ao_valv,	v_max_ao_valv, ar, m_valv, m_valv_kinetic, pg_m_valv, v_max_m_valv,	mr,	t_valv,	t_valv_kinetic,	pg_t_valv, v_max_t_valv, tr, others, pericard, cv, cv_colab, pa_press, nw_prognose, protocol]
                        data.to_excel('tabs/exams.xlsx', index=False)
                        report = docx.Document()
                        report.add_heading('', level=1)
                        report.add_paragraph(protocol_data)
                        report.save(report_filename)
                        st.success('Изменения в исследовании сохранены', icon="✅")
                    except:
                        st.error('Сохранение не произведено!', icon="⚠️")		                 
		
# КОНЕЦ ФУНКЦИИ АВТООБНОВЛЕНИЯ СТРАНИЦЫ
if __name__ == "__main__":
    run()
